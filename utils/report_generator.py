# -*- coding: utf-8 -*-
"""
utils/report_generator.py
~~~~~~~~~~~~~~~~~~~~~~~~~~
분기별/연간 성과보고서(.docx) 생성기. python-docx 기반.
Widget_Manager TaskStore 데이터 구조에 맞게 조정.
워크데이 계산은 utils.math_utils.count_workdays 재사용.
[의존성] requirements.txt에 python-docx 필요.
"""

import os
import datetime
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.math_utils import count_workdays

_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp")
_VIDEO_EXTS = (".mp4", ".mov", ".avi", ".mkv", ".webm")
_QUARTER_MONTHS = {1: (1, 3), 2: (4, 6), 3: (7, 9), 4: (10, 12)}
_MONTH_LASTDAY = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]


def period_range(year: int, mode: str, quarter: Optional[int] = None) -> Tuple[datetime.date, datetime.date]:
    if mode == "year":
        return datetime.date(year, 1, 1), datetime.date(year, 12, 31)
    if mode == "quarter":
        if quarter not in _QUARTER_MONTHS:
            raise ValueError("quarter는 1~4 사이여야 합니다.")
        sm, em = _QUARTER_MONTHS[quarter]
        last = _MONTH_LASTDAY[em - 1]
        if em == 2 and (year % 4 == 0 and (year % 100 != 0 or year % 400 == 0)):
            last = 29
        return datetime.date(year, sm, 1), datetime.date(year, em, last)
    raise ValueError("mode는 'year' 또는 'quarter' 여야 합니다.")


class ReportGenerator:
    def __init__(self, store):
        """store: TaskStore 인스턴스."""
        self._store = store

    def generate(
        self,
        year: int,
        mode: str,
        quarter: Optional[int] = None,
        author_name: str = "",
        holidays: Optional[List[str]] = None,
        output_path: Optional[str] = None,
    ) -> str:
        """보고서를 생성하고 저장 경로를 반환한다."""
        if holidays is None:
            holidays = []
        start, end = period_range(year, mode, quarter)
        tasks = self._collect_tasks(start, end)

        period_label = f"{year}년 연간" if mode == "year" else f"{year}년 {quarter}분기"

        if output_path is None:
            desktop = os.path.join(os.path.expanduser("~"), "Desktop")
            fname = f"성과보고서_{period_label.replace(' ', '_')}.docx"
            output_path = os.path.join(desktop, fname)

        doc = Document()
        self._add_title(doc, period_label, author_name, start, end)
        self._add_summary(doc, tasks, start, end, holidays)
        self._add_task_details(doc, tasks, holidays)
        doc.save(output_path)
        return output_path

    def _collect_tasks(self, start: datetime.date, end: datetime.date) -> List[dict]:
        """start~end 기간에 겹치는 일감 목록 반환 (시작일 순 정렬)."""
        tasks = self._store.by_date_range(start.isoformat(), end.isoformat())
        return sorted(tasks, key=lambda t: t.get("start", ""))

    def _add_title(self, doc, period_label, author_name, start, end):
        title = doc.add_heading(f"{period_label} 성과보고서", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(
            f"작성자: {author_name or '-'}    "
            f"보고 기간: {start.strftime('%Y.%m.%d')} ~ {end.strftime('%Y.%m.%d')}"
        )
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()

    def _add_summary(self, doc, tasks, start, end, holidays):
        doc.add_heading("요약 통계", level=1)
        total_tasks = len(tasks)
        done_tasks = sum(1 for t in tasks if t.get("status") == "done")
        total_wd = count_workdays(start, end, holidays)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        for label, value in [
            ("총 일감 수", f"{total_tasks}건"),
            ("완료 일감", f"{done_tasks}건"),
            ("기간 내 총 워크데이", f"{total_wd}일"),
            ("보고 기간", f"{start.strftime('%Y.%m.%d')} ~ {end.strftime('%Y.%m.%d')}"),
        ]:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
        doc.add_paragraph()

    def _add_task_details(self, doc, tasks, holidays):
        doc.add_heading("일감 상세", level=1)
        if not tasks:
            doc.add_paragraph("해당 기간에 등록된 일감이 없습니다.")
            return

        status_map = {"todo": "예정", "doing": "진행중", "done": "완료"}
        priority_map = {"high": "높음", "mid": "중간", "low": "낮음"}

        for task in tasks:
            try:
                task_start = datetime.date.fromisoformat(task["start"])
                task_end = datetime.date.fromisoformat(task["end"])
            except (KeyError, ValueError):
                continue

            wd = count_workdays(task_start, task_end, holidays)
            status_str = status_map.get(task.get("status", "todo"), "")
            priority_str = priority_map.get(task.get("priority", "mid"), "")

            doc.add_heading(task.get("title", "(제목 없음)"), level=2)
            meta = doc.add_paragraph()
            mrun = meta.add_run(
                f"기간: {task_start.strftime('%Y.%m.%d')} ~ {task_end.strftime('%Y.%m.%d')}  "
                f"({wd} 워크데이)  상태: {status_str}  우선순위: {priority_str}"
            )
            mrun.font.size = Pt(9)
            mrun.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            memo = (task.get("memo") or "").strip()
            if memo:
                doc.add_paragraph("메모", style="Intense Quote")
                for line in memo.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip(), style="List Bullet")

            jiras = task.get("jiras", [])
            folders = task.get("folders", [])
            if jiras or folders:
                doc.add_paragraph("링크 / 경로", style="Intense Quote")
                for j in jiras:
                    doc.add_paragraph(f"Jira: {j.get('name', j.get('path', ''))}", style="List Bullet")
                for f in folders:
                    doc.add_paragraph(f"폴더: {f.get('name', f.get('path', ''))}", style="List Bullet")

            attachments = task.get("attachments", [])
            imgs = [a for a in attachments if isinstance(a, str) and a.lower().endswith(_IMAGE_EXTS)]
            others = [a for a in attachments if a not in imgs]
            if imgs:
                doc.add_paragraph("첨부 이미지", style="Intense Quote")
                for a in imgs:
                    if os.path.exists(a):
                        try:
                            doc.add_picture(a, width=Inches(4.5))
                        except Exception:
                            doc.add_paragraph(f"(이미지 삽입 실패: {os.path.basename(a)})")
                    else:
                        doc.add_paragraph(f"(파일 없음: {os.path.basename(a)})")
            if others:
                doc.add_paragraph("첨부 파일 (경로)", style="Intense Quote")
                for a in others:
                    note = " — 동영상" if a.lower().endswith(_VIDEO_EXTS) else ""
                    doc.add_paragraph(f"{a}{note}", style="List Bullet")
            doc.add_paragraph()
